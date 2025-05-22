#src/vector_storage.py
import os
import asyncio
import asyncpg
import numpy as np
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from pathlib import Path
import PyPDF2
import docx
from mistralai import Mistral  # Updated import for new client
import logging
from dotenv import load_dotenv
from advanced_chunking import AdvancedTextChunker

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """Represents a chunk of text from a document"""
    content: str
    metadata: Dict[str, Any]
    page_number: Optional[int] = None
    chunk_index: Optional[int] = None

class DocumentProcessor:
    """Handles different document types and extracts text"""
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> List[DocumentChunk]:
        """Extract text from TXT file"""
        chunks = []
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read().strip()
                if content:
                    chunks.append(DocumentChunk(
                        content=content,
                        metadata={"source": file_path}
                    ))
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
        return chunks
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> List[DocumentChunk]:
        """Extract text from PDF file"""
        chunks = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        chunks.append(DocumentChunk(
                            content=text.strip(),
                            metadata={"source": file_path, "page": page_num + 1},
                            page_number=page_num + 1
                        ))
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
        return chunks
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> List[DocumentChunk]:
        """Extract text from DOCX file"""
        chunks = []
        try:
            doc = docx.Document(file_path)
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            if full_text:
                chunks.append(DocumentChunk(
                    content="\n".join(full_text),
                    metadata={"source": file_path}
                ))
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
        return chunks

class MistralEmbedder:
    """Handles embedding generation using Mistral AI (New Client)"""
    
    def __init__(self, api_key: str):
        self.client = Mistral(api_key=api_key)  # Updated for new client
        self.model = "mistral-embed"
    
    async def embed_text(self, text: str) -> List[float]:
        """Generate embedding for a single text"""
        try:
            response = self.client.embeddings.create(  # Updated method call
                model=self.model,
                inputs=[text]
            )
            return response.data[0].embedding
        except Exception as e:
            logger.error(f"Error generating embedding: {e}")
            return []
    
    async def embed_batch(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for multiple texts"""
        embeddings = []
        # Process in batches to avoid rate limits
        batch_size = 10
        
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]
            try:
                response = self.client.embeddings.create(  # Updated method call
                    model=self.model,
                    inputs=batch
                )
                batch_embeddings = [item.embedding for item in response.data]
                embeddings.extend(batch_embeddings)
            except Exception as e:
                logger.error(f"Error generating batch embeddings: {e}")
                # Add empty embeddings for failed batch
                embeddings.extend([[] for _ in batch])
        
        return embeddings

class VectorStore:
    """Handles vector storage operations with PostgreSQL + pgvector"""
    
    def __init__(self, database_url: str):
        self.database_url = database_url
        self.pool = None
    
    async def connect(self):
        """Establish connection pool"""
        self.pool = await asyncpg.create_pool(self.database_url)
        logger.info("Connected to database")
    
    async def close(self):
        """Close connection pool"""
        if self.pool:
            await self.pool.close()
            logger.info("Database connection closed")
    
    async def store_material(self, course_id: int, topic_id: int, material_name: str, 
                           file_path: str, file_type: str, uploaded_by: int) -> int:
        """Store course material metadata"""
        async with self.pool.acquire() as conn:
            material_id = await conn.fetchval(
                """
                INSERT INTO course_materials 
                (course_id, topic_id, material_name, file_path, file_type, uploaded_by)
                VALUES ($1, $2, $3, $4, $5, $6)
                RETURNING id
                """,
                course_id, topic_id, material_name, file_path, file_type, uploaded_by
            )
            logger.info(f"Stored material with ID: {material_id}")
            return material_id
    
    async def store_embedding(self, material_id: int, embedding: List[float]) -> int:
        """Store embedding vector"""
        async with self.pool.acquire() as conn:
            embedding_id = await conn.fetchval(
                """
                INSERT INTO embeddings (material_id, embedding)
                VALUES ($1, $2)
                RETURNING id
                """,
                material_id, embedding
            )
            
            # Update course_materials with embedding_id
            await conn.execute(
                """
                UPDATE course_materials 
                SET embedding_id = $1 
                WHERE id = $2
                """,
                embedding_id, material_id
            )
            
            logger.info(f"Stored embedding with ID: {embedding_id}")
            return embedding_id
    
    async def store_chunk_embedding(self, material_id: int, embedding: List[float], 
                                   chunk_metadata: Dict[str, Any]) -> int:
        """Store embedding vector with chunk metadata"""
        async with self.pool.acquire() as conn:
            # First, check if we need to create the chunks table
            await self._ensure_chunks_table(conn)
            
            # Store the embedding
            embedding_id = await conn.fetchval(
                """
                INSERT INTO embeddings (material_id, embedding)
                VALUES ($1, $2)
                RETURNING id
                """,
                material_id, embedding
            )
            
            # Store chunk metadata
            await conn.execute(
                """
                INSERT INTO document_chunks 
                (material_id, embedding_id, chunk_index, chunk_type, content_preview, 
                 concepts, formulas, difficulty, importance_score, metadata)
                VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10)
                """,
                material_id,
                embedding_id, 
                chunk_metadata.get('chunk_index', 0),
                chunk_metadata.get('chunk_type', 'paragraph'),
                chunk_metadata.get('content_preview', '')[:500],  # First 500 chars
                chunk_metadata.get('concepts', []),
                chunk_metadata.get('formulas', []),
                chunk_metadata.get('difficulty', []),
                chunk_metadata.get('importance_score', 0.0),
                chunk_metadata
            )
            
            logger.info(f"Stored chunk embedding with ID: {embedding_id}")
            return embedding_id
    
    async def _ensure_chunks_table(self, conn):
        """Create chunks table if it doesn't exist"""
        await conn.execute("""
            CREATE TABLE IF NOT EXISTS document_chunks (
                id SERIAL PRIMARY KEY,
                material_id INTEGER REFERENCES course_materials(id) ON DELETE CASCADE,
                embedding_id INTEGER REFERENCES embeddings(id) ON DELETE CASCADE,
                chunk_index INTEGER NOT NULL,
                chunk_type TEXT NOT NULL,
                content_preview TEXT,
                concepts TEXT[],
                formulas TEXT[],
                difficulty TEXT[],
                importance_score FLOAT DEFAULT 0.0,
                metadata JSONB,
                created_at TIMESTAMP DEFAULT NOW()
            );
            
            CREATE INDEX IF NOT EXISTS idx_chunks_material_id ON document_chunks(material_id);
            CREATE INDEX IF NOT EXISTS idx_chunks_embedding_id ON document_chunks(embedding_id);
            CREATE INDEX IF NOT EXISTS idx_chunks_type ON document_chunks(chunk_type);
            CREATE INDEX IF NOT EXISTS idx_chunks_concepts ON document_chunks USING GIN(concepts);
        """)
    
    async def similarity_search(self, query_embedding: List[float], 
                              course_id: Optional[int] = None, 
                              topic_id: Optional[int] = None, 
                              chunk_type: Optional[str] = None,
                              min_importance: Optional[float] = None,
                              limit: int = 5) -> List[Dict[str, Any]]:
        """Enhanced similarity search with chunk-aware filtering"""
        async with self.pool.acquire() as conn:
            # Check if document_chunks table exists
            chunks_exists = await conn.fetchval("""
                SELECT EXISTS (
                    SELECT FROM information_schema.tables 
                    WHERE table_schema = 'public' 
                    AND table_name = 'document_chunks'
                )
            """)
            
            if chunks_exists:
                return await self._chunk_based_search(
                    conn, query_embedding, course_id, topic_id, 
                    chunk_type, min_importance, limit
                )
            else:
                return await self._legacy_search(
                    conn, query_embedding, course_id, topic_id, limit
                )
    
    async def _chunk_based_search(self, conn, query_embedding: List[float], 
                                course_id: Optional[int] = None, 
                                topic_id: Optional[int] = None, 
                                chunk_type: Optional[str] = None,
                                min_importance: Optional[float] = None,
                                limit: int = 5) -> List[Dict[str, Any]]:
        """Perform chunk-based similarity search"""
        # Build query with optional filters
        where_conditions = ["e.embedding IS NOT NULL"]
        params = [query_embedding]
        param_count = 1
        
        if course_id:
            param_count += 1
            where_conditions.append(f"cm.course_id = ${param_count}")
            params.append(course_id)
        
        if topic_id:
            param_count += 1
            where_conditions.append(f"cm.topic_id = ${param_count}")
            params.append(topic_id)
        
        if chunk_type:
            param_count += 1
            where_conditions.append(f"dc.chunk_type = ${param_count}")
            params.append(chunk_type)
        
        if min_importance:
            param_count += 1
            where_conditions.append(f"dc.importance_score >= ${param_count}")
            params.append(min_importance)
        
        where_clause = " AND ".join(where_conditions)
        
        query = f"""
            SELECT 
                cm.id as material_id,
                cm.material_name,
                cm.file_path,
                cm.file_type,
                c.course_name,
                t.topic_name,
                dc.chunk_index,
                dc.chunk_type,
                dc.content_preview,
                dc.concepts,
                dc.formulas,
                dc.difficulty,
                dc.importance_score,
                dc.metadata,
                e.embedding <=> $1 as distance
            FROM course_materials cm
            JOIN embeddings e ON cm.id = e.material_id
            JOIN document_chunks dc ON e.id = dc.embedding_id
            LEFT JOIN courses c ON cm.course_id = c.id
            LEFT JOIN topics t ON cm.topic_id = t.id
            WHERE {where_clause}
            ORDER BY 
                dc.importance_score DESC,
                e.embedding <=> $1
            LIMIT ${param_count + 1}
        """
        
        params.append(limit)
        results = await conn.fetch(query, *params)
        
        return [dict(row) for row in results]
    
    async def _legacy_search(self, conn, query_embedding: List[float], 
                           course_id: Optional[int] = None, 
                           topic_id: Optional[int] = None, 
                           limit: int = 5) -> List[Dict[str, Any]]:
        """Fallback to legacy search if chunks table doesn't exist"""
        where_conditions = ["e.embedding IS NOT NULL"]
        params = [query_embedding, limit]
        param_count = 2
        
        if course_id:
            param_count += 1
            where_conditions.append(f"cm.course_id = ${param_count}")
            params.append(course_id)
        
        if topic_id:
            param_count += 1
            where_conditions.append(f"cm.topic_id = ${param_count}")
            params.append(topic_id)
        
        where_clause = " AND ".join(where_conditions)
        
        query = f"""
            SELECT 
                cm.id as material_id,
                cm.material_name,
                cm.file_path,
                cm.file_type,
                c.course_name,
                t.topic_name,
                e.embedding <=> $1 as distance
            FROM course_materials cm
            JOIN embeddings e ON cm.embedding_id = e.id
            LEFT JOIN courses c ON cm.course_id = c.id
            LEFT JOIN topics t ON cm.topic_id = t.id
            WHERE {where_clause}
            ORDER BY e.embedding <=> $1
            LIMIT $2
        """
        
        results = await conn.fetch(query, *params)
        return [dict(row) for row in results]

class MaterialProcessor:
    """Main class that orchestrates the entire pipeline"""
    
    def __init__(self, database_url: str, mistral_api_key: str):
        self.vector_store = VectorStore(database_url)
        self.embedder = MistralEmbedder(mistral_api_key)
        self.document_processor = DocumentProcessor()
        self.text_chunker = AdvancedTextChunker()
    
    async def initialize(self):
        """Initialize connections"""
        await self.vector_store.connect()
    
    async def cleanup(self):
        """Cleanup connections"""
        await self.vector_store.close()
    
    async def process_and_store_material(self, file_path: str, course_id: int, 
                                       topic_id: int, uploaded_by: int) -> bool:
        """Process a course material file and store it with embeddings (simple version)"""
        try:
            file_path_obj = Path(file_path)
            file_name = file_path_obj.name
            file_type = file_path_obj.suffix.lower()
            
            logger.info(f"Processing file: {file_name}")
            
            # Extract text based on file type
            if file_type == '.txt':
                chunks = self.document_processor.extract_text_from_txt(file_path)
            elif file_type == '.pdf':
                chunks = self.document_processor.extract_text_from_pdf(file_path)
            elif file_type == '.docx':
                chunks = self.document_processor.extract_text_from_docx(file_path)
            else:
                logger.error(f"Unsupported file type: {file_type}")
                return False
            
            if not chunks:
                logger.warning(f"No text extracted from {file_name}")
                return False
            
            # Combine all chunks into one text (simple approach)
            full_text = "\n\n".join([chunk.content for chunk in chunks])
            
            # Generate embedding
            embedding = await self.embedder.embed_text(full_text)
            
            if not embedding:
                logger.error(f"Failed to generate embedding for {file_name}")
                return False
            
            # Store material metadata
            material_id = await self.vector_store.store_material(
                course_id=course_id,
                topic_id=topic_id,
                material_name=file_name,
                file_path=file_path,
                file_type=file_type,
                uploaded_by=uploaded_by
            )
            
            # Store embedding
            await self.vector_store.store_embedding(material_id, embedding)
            
            logger.info(f"Successfully processed and stored {file_name}")
            return True
            
        except Exception as e:
            logger.error(f"Error processing material {file_path}: {e}")
            return False
    
    async def process_and_store_material_advanced(self, file_path: str, course_id: int, 
                                                topic_id: int, uploaded_by: int) -> bool:
        """Process material with advanced chunking and store multiple embeddings"""
        try:
            file_path_obj = Path(file_path)
            file_name = file_path_obj.name
            file_type = file_path_obj.suffix.lower()
            
            logger.info(f"Processing file with advanced chunking: {file_name}")
            
            # Extract text based on file type
            if file_type == '.txt':
                doc_chunks = self.document_processor.extract_text_from_txt(file_path)
            elif file_type == '.pdf':
                doc_chunks = self.document_processor.extract_text_from_pdf(file_path)
            elif file_type == '.docx':
                doc_chunks = self.document_processor.extract_text_from_docx(file_path)
            else:
                logger.error(f"Unsupported file type: {file_type}")
                return False
            
            if not doc_chunks:
                logger.warning(f"No text extracted from {file_name}")
                return False
            
            # Combine text from document chunks
            full_text = "\n\n".join([chunk.content for chunk in doc_chunks])
            
            # Apply advanced chunking
            advanced_chunks = self.text_chunker.create_advanced_chunks(
                full_text, 
                metadata={
                    "source_file": file_name,
                    "file_type": file_type,
                    "course_id": course_id,
                    "topic_id": topic_id
                }
            )
            
            # Store material metadata
            material_id = await self.vector_store.store_material(
                course_id=course_id,
                topic_id=topic_id,
                material_name=file_name,
                file_path=file_path,
                file_type=file_type,
                uploaded_by=uploaded_by
            )
            
            # Generate and store embeddings for each chunk
            successful_chunks = 0
            for chunk in advanced_chunks:
                try:
                    # Generate embedding for chunk
                    embedding = await self.embedder.embed_text(chunk.content)
                    
                    if embedding:
                        # Store embedding with enhanced metadata
                        await self.vector_store.store_chunk_embedding(
                            material_id=material_id,
                            embedding=embedding,
                            chunk_metadata={
                                "chunk_index": chunk.chunk_index,
                                "chunk_type": chunk.chunk_type.value,
                                "content_preview": chunk.content[:500],
                                "concepts": chunk.concepts,
                                "formulas": chunk.formulas,
                                "difficulty": chunk.difficulty_indicators,
                                "importance_score": chunk.importance_score,
                                **chunk.metadata
                            }
                        )
                        successful_chunks += 1
                    
                except Exception as e:
                    logger.error(f"Error processing chunk {chunk.chunk_index}: {e}")
            
            logger.info(f"Successfully processed {successful_chunks}/{len(advanced_chunks)} chunks for {file_name}")
            return successful_chunks > 0
            
        except Exception as e:
            logger.error(f"Error processing material {file_path}: {e}")
            return False
    
    async def search_similar_materials(self, query: str, course_id: Optional[int] = None, 
                                     limit: int = 5) -> List[Dict[str, Any]]:
        """Search for similar materials based on a query"""
        try:
            # Generate embedding for the query
            query_embedding = await self.embedder.embed_text(query)
            
            if not query_embedding:
                logger.error("Failed to generate query embedding")
                return []
            
            # Perform similarity search
            results = await self.vector_store.similarity_search(
                query_embedding=query_embedding,
                course_id=course_id,
                limit=limit
            )
            
            return results
            
        except Exception as e:
            logger.error(f"Error searching materials: {e}")
            return []